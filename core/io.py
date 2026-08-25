from __future__ import annotations

import io
import math
import re
from pathlib import Path
from typing import Any

import pandas as pd
from openpyxl.utils import get_column_letter

from .transactions import recommend_transaction
from .utils import unit_label

MAX_FINANCIAL_ROWS = 2_000
MAX_TRANSACTION_ROWS = 5_000
MAX_EXCEL_SHEETS = 20
MAX_PDF_PAGES = 40


def parse_number(value: Any) -> float | None:
    if value is None:
        return None
    if isinstance(value, bool):
        return float(value)
    if isinstance(value, (int, float)):
        try:
            n = float(value)
            return n if math.isfinite(n) else None
        except (TypeError, ValueError):
            return None

    text = str(value).strip()
    if not text:
        return None
    if text in {"-", "—"}:
        return 0.0
    parenthesized = bool(re.fullmatch(r"\(.*\)", text))
    cleaned = re.sub(r"[(),\s]", "", text)
    cleaned = re.sub(r"[^0-9.+-]", "", cleaned)
    if not cleaned or not re.search(r"\d", cleaned):
        return None
    try:
        number = float(cleaned)
    except ValueError:
        return None
    if not math.isfinite(number):
        return None
    if parenthesized:
        number = -abs(number)
    return number


def detect_unit_multiplier(text: str) -> int:
    text = str(text or "")
    patterns = [
        (r"단위\s*[:：]?\s*억원|억원\s*단위", 100_000_000),
        (r"단위\s*[:：]?\s*백만원|백만원\s*단위", 1_000_000),
        (r"단위\s*[:：]?\s*천원|천원\s*단위", 1_000),
    ]
    for pattern, multiplier in patterns:
        if re.search(pattern, text, flags=re.I):
            return multiplier
    return 1


def infer_statement(name: str) -> str:
    n = str(name or "").lower()
    if re.search(r"현금흐름|cash", n):
        return "현금흐름표"
    if re.search(r"재무상태|대차|balance", n):
        return "재무상태표"
    if re.search(r"손익|포괄손익|income|profit", n):
        return "손익계산서"
    return "기타"


def _matrix_from_df(df: pd.DataFrame) -> list[list[Any]]:
    return df.where(pd.notna(df), "").values.tolist()


def read_tabular_file(filename: str, data: bytes) -> dict[str, pd.DataFrame]:
    ext = Path(filename).suffix.lower()
    bio = io.BytesIO(data)
    if ext == ".csv":
        last_error: Exception | None = None
        for encoding in ("utf-8-sig", "cp949", "euc-kr", "utf-8"):
            try:
                bio.seek(0)
                df = pd.read_csv(bio, header=None, dtype=object, encoding=encoding)
                return {"CSV": df}
            except Exception as exc:  # pragma: no cover - encoding fallbacks vary by input
                last_error = exc
        raise ValueError(f"CSV 파일을 읽지 못했습니다: {last_error}")
    if ext in {".xlsx", ".xls"}:
        bio.seek(0)
        book = pd.ExcelFile(bio)
        out: dict[str, pd.DataFrame] = {}
        for sheet in book.sheet_names[:MAX_EXCEL_SHEETS]:
            df = pd.read_excel(book, sheet_name=sheet, header=None, dtype=object)
            if len(df) > 1:
                out[sheet] = df
        if not out:
            raise ValueError("읽을 수 있는 시트가 없습니다.")
        return out
    raise ValueError("지원하지 않는 표 형식입니다.")


def detect_financial_mapping(df: pd.DataFrame) -> dict[str, Any]:
    matrix = _matrix_from_df(df)
    header_index = 0
    for i, row in enumerate(matrix[:20]):
        low = [str(x).lower() for x in row]
        if len(row) >= 3 and any(re.search(r"계정|과목|account|항목", x) for x in low):
            header_index = i
            break

    raw = [str(x).strip() for x in matrix[header_index]]
    low = [x.lower() for x in raw]
    account_col = next((i for i, x in enumerate(low) if re.search(r"계정|과목|account|항목", x)), 0)

    year_cols: list[dict[str, Any]] = []
    for i, header in enumerate(raw):
        match = re.search(r"20\d{2}", header)
        if match:
            year_cols.append({"i": i, "label": match.group(0)})

    explicit_prior = next((i for i, x in enumerate(low) if re.search(r"전기|전년|prior|previous|비교", x)), -1)
    explicit_current = next((i for i, x in enumerate(low) if re.search(r"당기|당년|current|현재", x)), -1)

    if explicit_prior >= 0 and explicit_current >= 0 and explicit_prior != explicit_current:
        prior_col, current_col = explicit_prior, explicit_current
        prior_label = raw[prior_col] or "전기"
        current_label = raw[current_col] or "당기"
    elif len(year_cols) >= 2:
        ordered = sorted(year_cols, key=lambda x: int(x["label"]))
        prior_col, current_col = ordered[-2]["i"], ordered[-1]["i"]
        prior_label, current_label = ordered[-2]["label"], ordered[-1]["label"]
    else:
        candidates = [i for i in range(len(raw)) if i != account_col]
        prior_col = candidates[0] if candidates else 1
        current_col = candidates[1] if len(candidates) > 1 else min(prior_col + 1, max(len(raw) - 1, 0))
        prior_label = raw[prior_col] if prior_col < len(raw) and raw[prior_col] else "전기"
        current_label = raw[current_col] if current_col < len(raw) and raw[current_col] else "당기"

    period_cols = sorted(year_cols, key=lambda x: int(x["label"])) if len(year_cols) >= 2 else []
    return {
        "header_index": header_index,
        "headers": raw,
        "account_col": account_col,
        "prior_col": prior_col,
        "current_col": current_col,
        "prior_label": prior_label,
        "current_label": current_label,
        "period_cols": period_cols,
    }


def build_financial_rows(
    df: pd.DataFrame,
    *,
    filename: str,
    sheet_name: str,
    statement: str,
    account_col: int,
    prior_col: int,
    current_col: int,
    header_index: int,
    multiplier: int = 1,
    method: str = "excel",
    period_cols: list[dict[str, Any]] | None = None,
    id_start: int = 1,
) -> pd.DataFrame:
    if len({account_col, prior_col, current_col}) != 3:
        raise ValueError("계정과목·전기·당기 열은 서로 달라야 합니다.")

    matrix = _matrix_from_df(df)
    rows: list[dict[str, Any]] = []
    period_cols = period_cols or []
    for ri in range(header_index + 1, min(len(matrix), header_index + 1 + MAX_FINANCIAL_ROWS)):
        arr = matrix[ri]
        account = str(arr[account_col] if account_col < len(arr) else "").strip()
        if not account or re.search(r"단위|unit", account, flags=re.I) or re.fullmatch(r"\d+", account):
            continue
        prior_raw = arr[prior_col] if prior_col < len(arr) else None
        current_raw = arr[current_col] if current_col < len(arr) else None
        prior = parse_number(prior_raw)
        current = parse_number(current_raw)
        if prior is None and current is None:
            continue
        prior = None if prior is None else prior * multiplier
        current = None if current is None else current * multiplier

        def cell(col: int) -> str:
            if method == "excel":
                return f"{get_column_letter(col + 1)}{ri + 1}"
            return f"행 {ri + 1} · 열 {col + 1}"

        series: dict[str, float] = {}
        period_cells: dict[str, str] = {}
        for pc in period_cols:
            col = int(pc["i"])
            value = parse_number(arr[col] if col < len(arr) else None)
            if value is not None:
                series[str(pc["label"])] = value * multiplier
                period_cells[str(pc["label"])] = cell(col)

        rows.append(
            {
                "id": id_start + len(rows),
                "account": account,
                "statement": statement,
                "prior": prior,
                "current": current,
                "included": True,
                "series": series if len(series) >= 3 else {},
                "source_file": filename,
                "source_sheet": sheet_name,
                "source_row": ri + 1,
                "source_page": None,
                "prior_cell": cell(prior_col),
                "current_cell": cell(current_col),
                "period_cells": period_cells,
                "unit": unit_label(multiplier),
                "multiplier": int(multiplier),
                "method": method,
                "mapping_confirmed": True,
                "value_confirmed": True,
                "user_adjusted": False,
            }
        )
    return pd.DataFrame(rows)


def _parse_document_lines(
    lines: list[dict[str, Any]], *, filename: str, statement: str, multiplier: int, method: str
) -> pd.DataFrame:
    rows: list[dict[str, Any]] = []
    number_re = re.compile(r"\(?-?\d[\d,]*(?:\.\d+)?\)?")
    for item in lines:
        line = str(item.get("text", ""))
        nums = number_re.findall(line)
        if len(nums) < 2:
            continue
        first = line.find(nums[0])
        account = line[:first].strip().rstrip("|:")
        if not account or len(account) >= 90 or account.isdigit():
            continue
        prior = parse_number(nums[-2])
        current = parse_number(nums[-1])
        rows.append(
            {
                "id": len(rows) + 1,
                "account": account,
                "statement": statement,
                "prior": None if prior is None else prior * multiplier,
                "current": None if current is None else current * multiplier,
                "included": True,
                "series": {},
                "source_file": filename,
                "source_sheet": "",
                "source_row": item.get("line_no"),
                "source_page": item.get("page"),
                "prior_cell": "추출값",
                "current_cell": "추출값",
                "period_cells": {},
                "unit": unit_label(multiplier),
                "multiplier": int(multiplier),
                "method": method,
                "mapping_confirmed": True,
                "value_confirmed": False,
                "user_adjusted": False,
            }
        )
        if len(rows) >= 400:
            break
    return pd.DataFrame(rows)


def parse_document(filename: str, data: bytes, *, statement: str = "기타", multiplier: int = 1) -> tuple[pd.DataFrame, str]:
    ext = Path(filename).suffix.lower()
    lines: list[dict[str, Any]] = []
    preview_parts: list[str] = []
    if ext == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(data))
        for page_index, page in enumerate(reader.pages[:MAX_PDF_PAGES], start=1):
            text = page.extract_text() or ""
            preview_parts.append(f"[{page_index}페이지]\n{text}")
            for line_no, line in enumerate(text.splitlines(), start=1):
                if line.strip():
                    lines.append({"text": line, "page": page_index, "line_no": line_no})
        method = "pdf"
    elif ext == ".docx":
        from docx import Document

        doc = Document(io.BytesIO(data))
        raw_lines: list[str] = []
        raw_lines.extend(p.text for p in doc.paragraphs if p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                joined = " | ".join(cell.text.strip() for cell in row.cells)
                if joined.strip(" |"):
                    raw_lines.append(joined)
        preview_parts.append("\n".join(raw_lines))
        lines = [{"text": line, "line_no": i + 1, "page": None} for i, line in enumerate(raw_lines)]
        method = "docx"
    else:
        raise ValueError("PDF 또는 DOCX만 문서 추출을 지원합니다.")

    preview = "\n\n".join(preview_parts)
    return _parse_document_lines(lines, filename=filename, statement=statement, multiplier=multiplier, method=method), preview


def detect_transaction_mapping(df: pd.DataFrame) -> dict[str, Any]:
    matrix = _matrix_from_df(df)
    header_index = 0
    for i, row in enumerate(matrix[:20]):
        low = [str(x).lower() for x in row]
        if any(re.search(r"일자|날짜|date|적요|거래처|내용|description|금액|amount", x) for x in low):
            header_index = i
            break
    headers = [str(x).strip() for x in matrix[header_index]]
    low = [x.lower() for x in headers]

    def find(pattern: str) -> int:
        return next((i for i, x in enumerate(low) if re.search(pattern, x)), -1)

    date_col = find(r"일자|날짜|거래일|승인일|date")
    desc_col = find(r"적요|거래처|가맹점|상호|내용|description|memo|내역")
    amount_col = find(r"금액|승인금액|출금액|지출|amount")
    account_col = find(r"계정과목|계정명|account")
    date_col = 0 if date_col < 0 else date_col
    desc_col = min(1, max(len(headers) - 1, 0)) if desc_col < 0 else desc_col
    amount_col = min(2, max(len(headers) - 1, 0)) if amount_col < 0 else amount_col
    return {
        "header_index": header_index,
        "headers": headers,
        "date_col": date_col,
        "desc_col": desc_col,
        "amount_col": amount_col,
        "account_col": account_col,
    }


def build_transactions(
    df: pd.DataFrame,
    *,
    filename: str,
    date_col: int,
    desc_col: int,
    amount_col: int,
    account_col: int = -1,
    header_index: int = 0,
    existing: pd.DataFrame | None = None,
    id_start: int = 1,
) -> pd.DataFrame:
    if len({date_col, desc_col, amount_col}) != 3:
        raise ValueError("일자·적요·금액 열은 서로 달라야 합니다.")
    matrix = _matrix_from_df(df)
    built: list[dict[str, Any]] = []
    existing_records = existing.to_dict("records") if existing is not None and not existing.empty else []
    for ri in range(header_index + 1, min(len(matrix), header_index + 1 + MAX_TRANSACTION_ROWS)):
        arr = matrix[ri]
        desc = str(arr[desc_col] if desc_col < len(arr) else "").strip()
        amount = parse_number(arr[amount_col] if amount_col < len(arr) else None)
        if not desc or amount is None:
            continue
        original = str(arr[account_col] if account_col >= 0 and account_col < len(arr) else "").strip()
        rec = recommend_transaction(desc, existing_records)
        built.append(
            {
                "id": id_start + len(built),
                "date": str(arr[date_col] if date_col < len(arr) else "").strip(),
                "desc": desc,
                "amount": amount,
                "original_account": original,
                "ai_account": rec["account"],
                "recommendation_score": rec["score"],
                "recommendation_source": rec["source"],
                "final_account": original or rec["account"],
                "status": "대기",
                "user_modified": bool(original and original != rec["account"]),
                "duplicate": False,
                "source_file": filename,
                "source_row": ri + 1,
            }
        )
    return pd.DataFrame(built)
